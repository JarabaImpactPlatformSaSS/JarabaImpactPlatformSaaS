#!/usr/bin/env python3
"""
Generate comprehensive Jaraba Impact Platform SaaS documentation as .docx
Target: Executives and technicians of Plataforma de Ecosistemas Digitales S.L.
Date: 2026-03-04
"""

import os
from datetime import date
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ─── Constants ───
BLUE = RGBColor(0x23, 0x3D, 0x63)      # Azul corporativo
ORANGE = RGBColor(0xFF, 0x8C, 0x42)    # Naranja impulso
TEAL = RGBColor(0x00, 0xA9, 0xA5)      # Verde innovación
DARK = RGBColor(0x1A, 0x1A, 0x2E)
GRAY = RGBColor(0x64, 0x74, 0x8B)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BG = "F0F4F8"
HEADER_BG = "233D63"
ALT_ROW = "F8FAFC"

OUTPUT_DIR = "/home/PED/JarabaImpactPlatformSaaS/docs/tecnicos"
FILENAME = "20260304-Jaraba_Impact_Platform_SaaS_Documento_Comprensivo_v1.docx"


def set_cell_shading(cell, color_hex):
    """Apply background shading to a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_styled_table(doc, headers, rows, col_widths=None):
    """Create a professionally styled table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = WHITE
                run.font.size = Pt(9)
                run.font.name = 'Calibri'
        set_cell_shading(cell, HEADER_BG)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, value in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(value)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8.5)
                    run.font.name = 'Calibri'
                    run.font.color.rgb = DARK
            if r_idx % 2 == 1:
                set_cell_shading(cell, ALT_ROW)

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    return table


def add_kpi_box(doc, title, value, subtitle=""):
    """Add a KPI highlight paragraph."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{title}: ")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = BLUE
    run = p.add_run(value)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = ORANGE
    if subtitle:
        run = p.add_run(f"  {subtitle}")
        run.font.size = Pt(9)
        run.font.color.rgb = GRAY


def build_document():
    doc = Document()

    # ─── Page Setup ───
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)

    # ─── Default Font ───
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)
    font.color.rgb = DARK

    # Configure heading styles
    for level, size, color in [(1, 18, BLUE), (2, 14, BLUE), (3, 12, TEAL), (4, 10.5, DARK)]:
        hs = doc.styles[f'Heading {level}']
        hs.font.name = 'Calibri'
        hs.font.size = Pt(size)
        hs.font.color.rgb = color
        hs.font.bold = True

    # ═══════════════════════════════════════════════════════════════════
    # COVER PAGE
    # ═══════════════════════════════════════════════════════════════════
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("JARABA IMPACT PLATFORM")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = BLUE
    run.font.name = 'Calibri'

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Plataforma SaaS Multi-Vertical de Ecosistemas Digitales")
    run.font.size = Pt(14)
    run.font.color.rgb = TEAL
    run.font.name = 'Calibri'

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("DOCUMENTO COMPRENSIVO")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = ORANGE
    run.font.name = 'Calibri'

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Arquitectura, Logica de Negocio, Servicios, Sistemas y Procedimientos")
    run.font.size = Pt(11)
    run.font.color.rgb = GRAY

    for _ in range(4):
        doc.add_paragraph()

    # Metadata box
    meta_items = [
        ("Empresa", "Plataforma de Ecosistemas Digitales S.L."),
        ("Producto", "Jaraba Impact Platform SaaS"),
        ("Version documento", "1.0.0"),
        ("Fecha", "4 de marzo de 2026"),
        ("Clasificacion", "CONFIDENCIAL - Uso Interno"),
        ("Audiencia", "Directivos y Equipo Tecnico"),
    ]
    table = doc.add_table(rows=len(meta_items), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (key, val) in enumerate(meta_items):
        table.rows[i].cells[0].text = key
        table.rows[i].cells[1].text = val
        for p in table.rows[i].cells[0].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = BLUE
        for p in table.rows[i].cells[1].paragraphs:
            for run in p.runs:
                run.font.size = Pt(9)
        table.rows[i].cells[0].width = Cm(4)
        table.rows[i].cells[1].width = Cm(10)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════
    # TABLE OF CONTENTS (Manual)
    # ═══════════════════════════════════════════════════════════════════
    doc.add_heading("INDICE", level=1)
    toc_items = [
        "1. Resumen Ejecutivo",
        "2. Identidad del Producto y Propuesta de Valor",
        "3. Modelo de Negocio y Estrategia Comercial",
        "4. Arquitectura Tecnica General",
        "5. Arquitectura Multi-Tenant",
        "6. Ecosistema de 10 Verticales",
        "7. Stack Tecnologico Completo",
        "8. Catalogo de 95 Modulos Custom",
        "9. Modelo de Datos y Entidades",
        "10. Stack de Inteligencia Artificial",
        "11. Page Builder (GrapesJS)",
        "12. Sistema de Theming y Diseno",
        "13. SEO, GEO y Descubribilidad por IA",
        "14. Experiencia de Usuario (UX)",
        "15. Seguridad, Compliance y Gobernanza",
        "16. Infraestructura y DevOps",
        "17. Planes SaaS y Estructura de Precios",
        "18. Analisis Financiero y Proyecciones",
        "19. Estrategia Go-To-Market",
        "20. Roadmap y Madurez del Producto",
        "Anexo A: Listado Completo de Modulos",
        "Anexo B: Matriz de Entidades por Modulo",
        "Anexo C: Reglas y Directrices del Proyecto",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(0)
        for run in p.runs:
            run.font.size = Pt(10)
            run.font.color.rgb = DARK

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════
    # 1. RESUMEN EJECUTIVO
    # ═══════════════════════════════════════════════════════════════════
    doc.add_heading("1. Resumen Ejecutivo", level=1)

    doc.add_paragraph(
        "Jaraba Impact Platform es una plataforma SaaS multi-vertical desarrollada por "
        "Plataforma de Ecosistemas Digitales S.L. que integra comercio electronico, "
        "inteligencia artificial nativa, gestion de empleabilidad, emprendimiento, "
        "servicios profesionales y mas en un unico ecosistema digital. La plataforma "
        "sirve a 10 verticales de negocio desde una unica instancia multi-tenant, "
        "ofreciendo a cada organizacion un entorno digital completo y personalizable."
    )

    doc.add_heading("Cifras Clave del Producto", level=3)

    add_kpi_box(doc, "Codigo fuente", "712.000+ lineas PHP", "produccion")
    add_kpi_box(doc, "Modulos custom", "95 modulos", "Drupal 11")
    add_kpi_box(doc, "Entidades de datos", "443 tipos", "Content + Config")
    add_kpi_box(doc, "Agentes IA", "11 agentes Gen 2", "Claude + Gemini + OpenAI")
    add_kpi_box(doc, "Verticales", "10 sectores", "con logica de negocio propia")
    add_kpi_box(doc, "Tests automatizados", "453 tests", "Unit + Kernel")
    add_kpi_box(doc, "Rutas API/Web", "25.000+", "en routing.yml")
    add_kpi_box(doc, "Nivel de madurez", "5.0 / 5.0", "Production-Ready")

    doc.add_paragraph()
    doc.add_paragraph(
        "La plataforma se encuentra en estado pre-revenue con toda la infraestructura tecnica "
        "lista para el lanzamiento comercial. El foco inmediato es el Phase 0 (Piloto Institucional) "
        "con 3-5 entidades PIIL como primeros clientes, expandiendo progresivamente a comercio local, "
        "agroalimentario y servicios profesionales."
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════
    # 2. IDENTIDAD DEL PRODUCTO
    # ═══════════════════════════════════════════════════════════════════
    doc.add_heading("2. Identidad del Producto y Propuesta de Valor", level=1)

    doc.add_heading("2.1 Vision", level=2)
    doc.add_paragraph(
        "\"La primera plataforma de comercio disenada para que la Inteligencia Artificial "
        "venda tus productos.\" Jaraba Impact Platform conecta productores locales, "
        "emprendedores, comercios y profesionales con consumidores a traves de ecosistemas "
        "digitales verticalizados, potenciados por IA generativa nativa."
    )

    doc.add_heading("2.2 Filosofia \"Sin Humo\"", level=2)
    doc.add_paragraph(
        "El proyecto se rige por el principio \"Sin Humo\": codigo limpio, practico, "
        "sin complejidad innecesaria. Cada funcionalidad debe resolver un problema real "
        "y verificable. Las metricas de marketing solo incluyen datos verificables del codebase. "
        "Las proyecciones financieras usan escenarios conservadores."
    )

    doc.add_heading("2.3 Problema que Resuelve", level=2)
    problems = [
        ("Productores locales aislados", "Ecosistema digital conectado con marketplace y logistica"),
        ("Sin presencia digital", "Tienda online + IA + SEO/GEO en menos de 15 minutos"),
        ("Operaciones manuales", "Automatizacion completa (pedidos, facturas, envios, marketing)"),
        ("Trazabilidad inexistente", "Certificados digitales con QR y blockchain"),
        ("Ventas limitadas", "GEO (Generative Engine Optimization) para descubribilidad por ChatGPT/Gemini"),
        ("Herramientas fragmentadas", "Plataforma todo-en-uno por vertical (no 10 SaaS diferentes)"),
    ]
    add_styled_table(doc,
        ["Problema", "Solucion Jaraba"],
        problems,
        [6, 10]
    )

    doc.add_heading("2.4 Propuesta de Valor Unica (USP)", level=2)
    doc.add_paragraph(
        "Jaraba es la unica plataforma que combina cuatro dimensiones defensibles:"
    )
    usp_items = [
        "Integracion multi-vertical real (10 sectores, no 1)",
        "IA vertical nativa (11 agentes especializados por sector, no chatbot generico)",
        "Triple modelo de negocio (B2G institucional + B2B marketplace + licenciamiento SaaS)",
        "ADN rural (baja conectividad, cumplimiento institucional PIIL/FSE/SEPE, accesibilidad)",
    ]
    for item in usp_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading("2.5 Segmentos de Mercado Objetivo", level=2)
    segments = [
        ("Entidades de empleo (PIIL, SAE, GDR)", "Empleabilidad", "79 EUR/mes"),
        ("Productores agroalimentarios", "AgroConecta", "39 EUR/mes"),
        ("Comercio local y retail", "ComercioConecta", "29 EUR/mes"),
        ("Profesionales de servicios", "ServiciosConecta", "79 EUR/mes"),
        ("Cooperativas y asociaciones", "AgroConecta/Emprendimiento", "149 EUR/mes"),
        ("Instituciones publicas", "White-label multi-vertical", "Licencia anual"),
        ("Despachos de abogados", "JarabaLex", "99 EUR/mes"),
        ("Centros de formacion", "Formacion", "79 EUR/mes"),
    ]
    add_styled_table(doc,
        ["Segmento", "Vertical Principal", "Precio Orientativo"],
        segments,
        [6, 4, 4]
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════
    # 3. MODELO DE NEGOCIO
    # ═══════════════════════════════════════════════════════════════════
    doc.add_heading("3. Modelo de Negocio y Estrategia Comercial", level=1)

    doc.add_heading("3.1 Triple Modelo de Ingresos", level=2)
    revenue_models = [
        ("B2G Institucional", "Licencias anuales a ayuntamientos, diputaciones, GDR, PIIL",
         "White-label, multi-programa, cumplimiento FSE/SEPE, dashboards de impacto"),
        ("B2B SaaS", "Suscripciones mensuales/anuales a empresas y profesionales",
         "3 planes (Starter/Professional/Enterprise), self-service + Kit Digital"),
        ("Marketplace", "Comision por transaccion en marketplace (8%)",
         "Stripe Connect destination charges, split automatico"),
    ]
    add_styled_table(doc,
        ["Canal", "Descripcion", "Mecanismo"],
        revenue_models,
        [3, 5.5, 7]
    )

    doc.add_heading("3.2 Planes de Suscripcion SaaS", level=2)
    plans = [
        ("Starter", "0 EUR", "0 EUR", "Funcionalidades basicas, 1 usuario, sin IA, ideal para probar"),
        ("Professional", "29 EUR", "290 EUR", "IA incluida, hasta 5 usuarios, FNMT, analytics"),
        ("Enterprise", "99 EUR", "990 EUR", "Ilimitado, white-label, blockchain, soporte premium"),
    ]
    add_styled_table(doc,
        ["Plan", "Precio/Mes", "Precio/Ano", "Incluye"],
        plans,
        [2.5, 2.5, 2.5, 8]
    )

    doc.add_paragraph()
    doc.add_paragraph(
        "Los precios se verticalizan por sector: cada vertical puede tener limites, "
        "features y precios diferentes segun su SaasPlan (ContentEntity) y SaasPlanTier "
        "(ConfigEntity). El sistema de cascada PlanResolverService busca primero "
        "{vertical}_{tier}, luego _default_{tier}, y finalmente datos fallback hardcoded."
    )

    doc.add_heading("3.3 Estrategia Kit Digital", level=2)
    doc.add_paragraph(
        "Espana ha emitido 880.000+ bonos Kit Digital (2022-2025) de entre 2.000 y 12.000 EUR "
        "para digitalizacion de pymes. Jaraba se posiciona como solucion 100% subvencionable: "
        "el plan Starter anual (348 EUR) queda cubierto integramente por el bono minimo de 3.000 EUR. "
        "Esto reduce la friccion de adopcion a cero coste para el cliente."
    )

    doc.add_heading("3.4 Ventaja Competitiva vs Alternativas", level=2)
    competitors = [
        ("Shopify", "Solo e-commerce", "36 EUR/mes", "No IA nativa, no multi-vertical, no institucional"),
        ("Holded", "Solo contabilidad/ERP", "29 EUR/mes", "No marketplace, no IA, no page builder"),
        ("Factorial", "Solo RRHH", "4,5 EUR/empleado", "Solo un vertical, no produccion, no comercio"),
        ("Agroptima", "Solo gestion agraria", "20 EUR/mes", "Solo parcelas, no marketplace, no trazabilidad QR"),
        ("Jaraba", "Multi-vertical completo", "0-99 EUR/mes", "10 verticales + IA + marketplace + compliance"),
    ]
    add_styled_table(doc,
        ["Plataforma", "Alcance", "Precio", "Limitaciones vs Jaraba"],
        competitors,
        [2.5, 3, 3, 7]
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════
    # 4. ARQUITECTURA TECNICA GENERAL
    # ═══════════════════════════════════════════════════════════════════
    doc.add_heading("4. Arquitectura Tecnica General", level=1)

    doc.add_heading("4.1 Diagrama de Capas", level=2)
    doc.add_paragraph(
        "La arquitectura sigue un modelo de capas con separacion clara de responsabilidades:"
    )

    layers = [
        ("Capa de Presentacion", "Twig templates, SCSS/CSS, Vanilla JS, GrapesJS Page Builder\n"
         "Zero Region Pattern: clean_content + drupalSettings via preprocess"),
        ("Capa de Aplicacion", "95 modulos custom Drupal 11, 443 entity types, 25.000+ rutas\n"
         "PremiumEntityFormBase, Slide-Panel Forms, Controllers con DI"),
        ("Capa de Servicios", "149+ services en core, TenantBridge, ModelRouter, Streaming\n"
         "Inyeccion de dependencias, optional services @?, tagged services"),
        ("Capa de IA", "11 agentes Gen 2 (SmartBaseAgent), MCP Server, RAG/Qdrant\n"
         "Streaming SSE, PII Guardrails, Semantic Cache, Tool Registry"),
        ("Capa de Datos", "MariaDB 10.11 (relacional), Redis 7.4 (cache/queue)\n"
         "Qdrant (vector DB), Apache Tika (document parsing)"),
        ("Capa de Infraestructura", "IONOS Dedicated L-16 NVMe, GitHub Actions CI/CD\n"
         "8 workflows, deploy automatico, security scan diario"),
    ]
    add_styled_table(doc,
        ["Capa", "Componentes"],
        layers,
        [4, 12]
    )

    doc.add_heading("4.2 Principios Arquitecturales", level=2)
    principles = [
        "Single-Instance Multi-Tenant: Una instalacion Drupal sirve a todos los tenants",
        "Soft Isolation via Group Module: Aislamiento logico de contenido por grupo",
        "Configuration-Driven: Verticales y planes como ConfigEntities, no hardcoded",
        "AI-Native: IA integrada en el core, no como addon externo",
        "API-First: Todos los servicios accesibles via API REST/JSON-RPC",
        "Compliance-by-Design: GDPR, LOPD, ENS, SOC2, ISO27001 desde el diseno",
        "Progressive Enhancement: Funcionalidad base sin JS, mejora con JS",
        "Defensive Programming: Fallbacks en cascada, circuit breakers, self-healing",
    ]
    for p_text in principles:
        doc.add_paragraph(p_text, style='List Bullet')

    doc.add_heading("4.3 Patrones de Diseno Clave", level=2)
    patterns = [
        ("Cascade Resolution", "Precio/features busca: vertical_tier -> default_tier -> fallback"),
        ("Bridge Pattern", "TenantBridgeService mapea Tenant (billing) <-> Group (content)"),
        ("Strategy + Registry", "ToolRegistry con tagged services para herramientas IA"),
        ("Circuit Breaker", "ProviderFallbackService con estados open/closed/half-open"),
        ("Observer/Event", "Hooks Drupal + Event Subscribers para desacoplamiento"),
        ("State Machine", "SupportTicket con 10 estados (new -> closed) y transiciones definidas"),
        ("CQRS Ligero", "Entity Query para lectura, EntityDefinitionUpdateManager para escritura"),
        ("Semantic Cache", "Qdrant fuzzy matching antes de llamada LLM (reduce coste 40%)"),
    ]
    add_styled_table(doc,
        ["Patron", "Aplicacion en Jaraba"],
        patterns,
        [4, 12]
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════
    # 5. ARQUITECTURA MULTI-TENANT
    # ═══════════════════════════════════════════════════════════════════
    doc.add_heading("5. Arquitectura Multi-Tenant", level=1)

    doc.add_heading("5.1 Modelo de Aislamiento", level=2)
    doc.add_paragraph(
        "Jaraba utiliza un modelo Single-Instance Multi-Tenant con soft isolation via "
        "el modulo Group de Drupal. Esto significa que una unica instalacion de Drupal "
        "sirve a multiples organizaciones (tenants), cada una con su propio espacio de "
        "contenido, usuarios, configuracion visual y plan de suscripcion."
    )

    doc.add_heading("5.2 Jerarquia de Entidades", level=2)
    hierarchy = [
        ("Vertical", "ConfigEntity", "Define el sector (empleabilidad, agroconecta, etc.)"),
        ("Tenant", "ContentEntity", "Entidad de facturacion: plan, Stripe, dominio"),
        ("Group", "ContentEntity (contrib)", "Aislamiento de contenido: usuarios, paginas, productos"),
        ("User", "ContentEntity (core)", "Usuario final con roles y tenant_id"),
        ("SaasPlan", "ContentEntity", "Precios y limites por vertical+tier"),
        ("SaasPlanTier", "ConfigEntity", "Definicion de tier (starter/professional/enterprise)"),
    ]
    add_styled_table(doc,
        ["Entidad", "Tipo", "Funcion"],
        hierarchy,
        [3, 3.5, 9]
    )

    doc.add_heading("5.3 Servicios de Multi-Tenancy", level=2)
    mt_services = [
        ("TenantContextService", "ecosistema_jaraba_core.tenant_context",
         "Resuelve el tenant actual del usuario via admin_user + group membership"),
        ("TenantBridgeService", "ecosistema_jaraba_core.tenant_bridge",
         "Mapeo bidireccional Tenant <-> Group. OBLIGATORIO para toda resolucion"),
        ("TenantResolverService", "ecosistema_jaraba_core.tenant_resolver",
         "getCurrentTenant() devuelve GroupInterface (no TenantInterface)"),
        ("QuotaManagerService", "jaraba_page_builder.quota_manager",
         "Verifica limites del plan (paginas, storage, usuarios, API calls)"),
        ("ThemeTokenService", "ecosistema_jaraba_core.theme_token",
         "Cascada de tokens visuales: base -> vertical -> tenant override"),
    ]
    add_styled_table(doc,
        ["Servicio", "Service ID", "Funcion"],
        mt_services,
        [3.5, 5, 7]
    )

    doc.add_heading("5.4 Reglas Cardinales de Tenancy", level=2)
    rules = [
        "TENANT-001: TODA query a base de datos DEBE filtrar por tenant_id. Sin excepciones.",
        "TENANT-002: Usar ecosistema_jaraba_core.tenant_context para obtener tenant. NUNCA via queries ad-hoc.",
        "TENANT-BRIDGE-001: SIEMPRE usar TenantBridgeService para resolver Tenant<->Group.",
        "TENANT-ISOLATION-ACCESS-001: Todo AccessControlHandler con tenant_id DEBE verificar tenant match.",
        "ACCESS-STRICT-001: Comparaciones de ownership con (int) === (int), NUNCA ==.",
    ]
    for rule in rules:
        doc.add_paragraph(rule, style='List Bullet')

    doc.add_heading("5.5 Estrategia de Escalado", level=2)
    scaling = [
        ("Fase 1: Single Server", "0-50 tenants", "Servidor actual IONOS L-16 (128GB RAM)",
         "Suficiente para Phase 0-1 del GTM"),
        ("Fase 2: DB Separada", "50-200 tenants", "Servidor de aplicacion + DB dedicada",
         "MariaDB en instancia separada, Redis cluster"),
        ("Fase 3: Load Balanced", "200+ tenants", "Multiples app servers + DB cluster",
         "HAProxy/Nginx LB, MariaDB Galera, Redis Sentinel"),
    ]
    add_styled_table(doc,
        ["Fase", "Tenants", "Infraestructura", "Notas"],
        scaling,
        [3, 2.5, 5, 5]
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════
    # 6. ECOSISTEMA DE 10 VERTICALES
    # ═══════════════════════════════════════════════════════════════════
    doc.add_heading("6. Ecosistema de 10 Verticales", level=1)

    doc.add_paragraph(
        "Cada vertical representa un sector de mercado con logica de negocio, entidades, "
        "servicios y agentes IA especializados. Los verticales comparten infraestructura "
        "pero ofrecen experiencias completamente diferenciadas."
    )

    verticals = [
        ("empleabilidad", "Empleabilidad",
         "Bolsa de empleo, LMS, matching candidato-oferta, CV builder, diagnostico, "
         "copilot para buscadores de empleo. Cumplimiento SEPE/PIIL/FSE.",
         "14 entidades, 12+ servicios, 6 modulos",
         "PIIL, SAE, GDR, Agencias de empleo"),
        ("emprendimiento", "Emprendimiento",
         "Business Canvas con 44 experimentos Osterwalder, funding tracker, red de mentores, "
         "diagnostico express, copilot para emprendedores.",
         "10 entidades, 10+ servicios, 6 modulos",
         "Emprendedores, incubadoras, aceleradoras"),
        ("comercioconecta", "ComercioConecta",
         "E-commerce omnicanal, POS sync, QR dinamicos, ofertas flash, fidelizacion, "
         "marketplace de proximidad.",
         "42 entidades, 25+ servicios, 97 rutas",
         "Comercio minorista, tiendas locales"),
        ("agroconecta", "AgroConecta",
         "Marketplace agroalimentario, trazabilidad QR, certificados digitales FNMT, "
         "portal productor/consumidor, blockchain.",
         "20 entidades, 18+ servicios",
         "Productores, cooperativas, DO/IGP"),
        ("jarabalex", "JarabaLex",
         "Gestion de casos legales, calendario de plazos, boveda digital, templates, "
         "facturacion legal, LexNET, investigacion IA.",
         "7+ entidades, 15+ servicios, 12 submodulos",
         "Abogados, despachos, asesorias"),
        ("serviciosconecta", "ServiciosConecta",
         "Motor de reservas, calendario disponibilidad, firma digital PAdES, "
         "presupuestos IA, triaje inteligente.",
         "6 entidades, 5+ servicios",
         "Profesionales freelance, consultores"),
        ("andalucia_ei", "Andalucia +ei",
         "Programa regional de emprendimiento: gestion participantes, tracking impacto, "
         "reporting institucional.",
         "Programa especifico",
         "Entidades Andalucia Emprende"),
        ("jaraba_content_hub", "Content Hub",
         "Blog con IA, gestion de autores editoriales, SEO avanzado, RSS 2.0, "
         "newsletter, comentarios, categorias.",
         "5+ entidades, 16 servicios",
         "Creadores de contenido, marketing"),
        ("formacion", "Formacion",
         "LMS con cursos, certificaciones, contenido interactivo H5P, "
         "rutas de aprendizaje personalizadas por IA.",
         "Integrado con LMS",
         "Centros formativos, e-learning"),
        ("demo", "Demo",
         "Sandbox con datos de ejemplo para demostraciones, tours guiados, "
         "producto interactivo para ventas.",
         "Configuracion especial",
         "Prospects, equipo comercial"),
    ]

    for key, name, desc, stats, audience in verticals:
        doc.add_heading(f"6.{verticals.index((key, name, desc, stats, audience))+1} {name} ({key})", level=3)
        doc.add_paragraph(desc)
        p = doc.add_paragraph()
        run = p.add_run("Alcance tecnico: ")
        run.bold = True
        run.font.size = Pt(9)
        p.add_run(stats).font.size = Pt(9)
        p = doc.add_paragraph()
        run = p.add_run("Audiencia: ")
        run.bold = True
        run.font.size = Pt(9)
        p.add_run(audience).font.size = Pt(9)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════
    # 7. STACK TECNOLOGICO
    # ═══════════════════════════════════════════════════════════════════
    doc.add_heading("7. Stack Tecnologico Completo", level=1)

    doc.add_heading("7.1 Backend", level=2)
    backend = [
        ("Framework CMS", "Drupal 11.x", "Core del ecosistema, entity system, routing, DI"),
        ("Lenguaje", "PHP 8.4+", "strict_types, readonly properties, typed properties"),
        ("Base de datos", "MariaDB 10.11", "159+ tablas por tenant, _field_data para translatables"),
        ("Cache / Queue", "Redis 7.4", "Cache backend, queue worker, session storage"),
        ("Vector DB", "Qdrant 1.16", "Embeddings para RAG, semantic cache, long-term memory"),
        ("Document Parsing", "Apache Tika", "Extraccion de texto para indexacion en knowledge base"),
        ("Package Manager", "Composer 2.x", "40+ dependencias PHP, PSR-4 autoloading"),
        ("CLI", "Drush 13", "Administracion, migraciones, cache, cron, deploy hooks"),
    ]
    add_styled_table(doc,
        ["Componente", "Version", "Uso"],
        backend,
        [3, 3, 10]
    )

    doc.add_heading("7.2 Frontend", level=2)
    frontend = [
        ("Templates", "Twig (Drupal native)", "65+ parciales, Zero Region Pattern, {% trans %}"),
        ("Estilos", "SCSS / Dart Sass", "102 archivos SCSS, 5 capas de tokens, --ej-* vars"),
        ("JavaScript", "Vanilla JS", "Drupal.behaviors, NO frameworks (React/Vue/Angular)"),
        ("Page Builder", "GrapesJS 5.7", "202+ bloques, 24 categorias, 11 plugins custom"),
        ("Iconos", "SVG inline", "352 iconos, 6 categorias, variantes outline/duotone"),
        ("Tipografia", "Outfit (Google)", "Font primaria de marca, BRAND-FONT-001"),
    ]
    add_styled_table(doc,
        ["Componente", "Tecnologia", "Detalle"],
        frontend,
        [3, 3.5, 9]
    )

    doc.add_heading("7.3 Integraciones Externas", level=2)
    integrations = [
        ("Pagos", "Stripe Connect", "Split payments, suscripciones, webhooks HMAC"),
        ("IA - LLM", "Claude API (Anthropic)", "Agentes Gen 2, razonamiento avanzado"),
        ("IA - LLM", "Gemini API (Google)", "Tier fast, grounding estricto"),
        ("IA - LLM", "OpenAI API", "GPT-4o Vision, Whisper, DALL-E 3, TTS"),
        ("Email", "SMTP (configurable)", "28 templates MJML, CAN-SPAM compliance"),
        ("Auth Social", "Google, LinkedIn, Microsoft", "OAuth 2.0 via social_auth"),
        ("Firma Digital", "AutoFirma / FNMT", "Certificados digitales PAdES"),
        ("Mapas", "Leaflet + Geocoder", "Geolocalizacion, mapas interactivos"),
        ("Analytics", "GTM + custom", "Heatmaps, A/B testing, pixel manager"),
    ]
    add_styled_table(doc,
        ["Area", "Tecnologia", "Detalle"],
        integrations,
        [2.5, 3.5, 10]
    )

    doc.add_heading("7.4 Dependencias Composer Principales", level=2)
    deps = [
        ("drupal/commerce", "^3.1", "E-commerce framework completo"),
        ("drupal/group", "^3.2", "Multi-tenancy content isolation"),
        ("drupal/domain", "^2.0", "Custom domains per tenant"),
        ("drupal/ai", "latest", "AI provider abstraction layer"),
        ("stripe/stripe-php", "^15.0", "Stripe API client"),
        ("drupal/search_api", "latest", "Search abstraction (Solr/DB)"),
        ("dompdf/dompdf", "^2.0", "Generacion de PDFs (facturas, certificados)"),
        ("tecnickcom/tcpdf", "latest", "PDF avanzado (firmas digitales)"),
        ("drupal/eca", "^3.0", "Event-Condition-Action automation"),
        ("defuse/php-encryption", "latest", "Cifrado de datos sensibles"),
    ]
    add_styled_table(doc,
        ["Paquete", "Version", "Proposito"],
        deps,
        [4, 2, 10]
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════
    # 8. CATALOGO DE MODULOS
    # ═══════════════════════════════════════════════════════════════════
    doc.add_heading("8. Catalogo de 95 Modulos Custom", level=1)

    doc.add_paragraph(
        "Los 95 modulos custom se organizan en 10 tiers funcionales. Todos siguen "
        "el prefijo jaraba_* (excepto ecosistema_jaraba_core y ai_provider_google_gemini)."
    )

    doc.add_heading("8.1 Tier 1: Infraestructura Core (3 modulos)", level=2)
    core_mods = [
        ("ecosistema_jaraba_core", "149+ services, 35 entities",
         "Fundacion multi-tenant: billing, compliance, firma digital, tenant mgmt, "
         "pricing cascade, onboarding, API keys, webhooks, search"),
        ("jaraba_theming", "Design tokens, presets",
         "Tokens CSS, personalizacion visual multi-tenant, presets por vertical"),
        ("ai_provider_google_gemini", "AI provider",
         "Integracion con Google Gemini API para tier fast"),
    ]
    add_styled_table(doc,
        ["Modulo", "Alcance", "Descripcion"],
        core_mods,
        [4.5, 3.5, 8]
    )

    doc.add_heading("8.2 Tier 2: Stack de IA (4 modulos)", level=2)
    ai_mods = [
        ("jaraba_ai_agents", "56 services, 20 entities",
         "11 agentes Gen 2, tool registry, brand voice, observability, guardrails, MCP server"),
        ("jaraba_copilot_v2", "27 services, 10 entities",
         "Copilot de emprendimiento: 5 modos, 44 experimentos Osterwalder, semantic cache, streaming"),
        ("jaraba_rag", "7 services",
         "RAG pipeline: Qdrant vector DB, semantic search, document chunking, re-ranking"),
        ("jaraba_agents", "15 services",
         "Orquestacion autonoma de agentes, aprobaciones, long-term memory, handoff decisions"),
    ]
    add_styled_table(doc,
        ["Modulo", "Alcance", "Descripcion"],
        ai_mods,
        [4, 3.5, 8.5]
    )

    doc.add_heading("8.3 Tier 3: Contenido y Page Builder (3 modulos)", level=2)
    content_mods = [
        ("jaraba_page_builder", "24 services, 9 entities",
         "Landing pages con GrapesJS, 202+ bloques, quotas, templates, accesibilidad, SEO"),
        ("jaraba_content_hub", "16 services, 5+ entities",
         "Blog AI, articulos, autores editoriales, SEO, RSS 2.0, categorias, comentarios"),
        ("jaraba_site_builder", "9 entities",
         "Construccion de sitios multi-pagina completos"),
    ]
    add_styled_table(doc,
        ["Modulo", "Alcance", "Descripcion"],
        content_mods,
        [4, 3.5, 8.5]
    )

    doc.add_heading("8.4 Tier 4: Soporte y CRM (3 modulos)", level=2)
    support_mods = [
        ("jaraba_support", "20 services, 9 entities",
         "Tickets world-class, clasificacion IA, SLA engine, maquina de estados 10 estados"),
        ("jaraba_crm", "CRM completo",
         "Pipeline de ventas, contactos, empresas, actividades, scoring"),
        ("jaraba_customer_success", "9 entities",
         "Prediccion de churn, health scoring, playbooks de retencion, NPS"),
    ]
    add_styled_table(doc,
        ["Modulo", "Alcance", "Descripcion"],
        support_mods,
        [4.5, 3, 8.5]
    )

    doc.add_heading("8.5 Tier 5: Billing y Comercio (4 modulos)", level=2)
    billing_mods = [
        ("jaraba_billing", "20+ services, 5 entities",
         "Suscripciones Stripe, metering, trials, impact credits, facturas"),
        ("jaraba_commerce", "Commerce extension",
         "Extension de Drupal Commerce 3.x para multi-tenant marketplace"),
        ("jaraba_usage_billing", "Pipeline avanzado",
         "Facturacion basada en uso (API calls, storage, AI tokens)"),
        ("jaraba_addons", "Add-ons",
         "Complementos por suscripcion (features extra, storage, API calls)"),
    ]
    add_styled_table(doc,
        ["Modulo", "Alcance", "Descripcion"],
        billing_mods,
        [4, 3, 9]
    )

    doc.add_heading("8.6 Tier 6: Modulos Verticales (30+ modulos)", level=2)
    doc.add_paragraph(
        "Cada vertical tiene entre 1 y 12 submodulos especializados. "
        "Ver Seccion 6 para detalle funcional por vertical. "
        "El listado completo se encuentra en el Anexo A."
    )

    doc.add_heading("8.7 Tier 7-10: Plataforma Avanzada (50+ modulos)", level=2)
    advanced_categories = [
        ("Analytics y Testing", "jaraba_analytics, jaraba_ab_testing, jaraba_heatmap, jaraba_pixels",
         "Tracking completo, A/B testing, heatmaps, pixel manager"),
        ("Comunicacion", "jaraba_email, jaraba_messaging, jaraba_notifications, jaraba_social",
         "28 templates MJML, mensajeria interna, push notifications, social media"),
        ("Compliance", "jaraba_privacy, jaraba_governance, jaraba_security_compliance",
         "GDPR/LOPD, SOC2, ISO27001, ENS, clasificacion de datos"),
        ("Facturacion Electronica", "jaraba_facturae, jaraba_verifactu, jaraba_einvoice_b2b",
         "Facturae 3.2.2, VeriFactu RD 1007/2023, UBL 2.1 + EN 16931"),
        ("Integraciones", "jaraba_integrations, jaraba_connector_sdk, jaraba_sso",
         "API marketplace, SAML 2.0, OIDC, SCIM 2.0"),
        ("White-label", "jaraba_whitelabel, jaraba_pwa, jaraba_mobile",
         "Dominios custom, PWA, backend para app nativa"),
        ("Credenciales", "jaraba_credentials, jaraba_identity, jaraba_zkp",
         "Open Badge 3.0, identidad soberana (DID), zero-knowledge proofs"),
        ("Institucional", "jaraba_institutional, jaraba_sepe_teleformacion",
         "Programas PIIL/FSE/FUNDAE, compliance SEPE, teleformacion"),
        ("Operaciones", "jaraba_dr, jaraba_sla, jaraba_foc, jaraba_performance",
         "Disaster recovery, SLA management, FinOps center, performance monitoring"),
    ]
    add_styled_table(doc,
        ["Categoria", "Modulos", "Funcionalidad"],
        advanced_categories,
        [3, 6, 7]
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════
    # 9. MODELO DE DATOS
    # ═══════════════════════════════════════════════════════════════════
    doc.add_heading("9. Modelo de Datos y Entidades", level=1)

    doc.add_heading("9.1 Resumen Cuantitativo", level=2)
    doc.add_paragraph(
        "El sistema define 443 tipos de entidad (ContentEntity + ConfigEntity) "
        "distribuidos en 95 modulos. Cada entidad sigue las convenciones del proyecto: "
        "AccessControlHandler obligatorio, tenant_id como entity_reference, "
        "PremiumEntityFormBase para formularios, Views integration."
    )

    doc.add_heading("9.2 Entidades Core", level=2)
    core_entities = [
        ("Tenant", "ContentEntity", "Facturacion: plan, Stripe customer, dominio, vertical"),
        ("Vertical", "ContentEntity (translatable)", "machine_name, status, descripcion, icono"),
        ("SaasPlan", "ContentEntity (translatable)", "Precios por vertical+tier (EUR mensual/anual)"),
        ("SaasPlanTier", "ConfigEntity", "Definicion de tier: starter/professional/enterprise"),
        ("Feature", "ContentEntity", "Features habilitables por plan/tenant"),
        ("FeatureFlag", "ContentEntity", "Feature flags para activacion gradual"),
        ("FreemiumVerticalLimit", "ContentEntity", "Limites del plan gratuito por vertical"),
        ("Badge / BadgeAward", "ContentEntity", "Gamificacion: insignias y otorgamientos"),
        ("AuditLog", "ContentEntity", "Log inmutable de acciones (append-only)"),
        ("DesignTokenConfig", "ContentEntity", "Tokens de diseno por tenant"),
        ("PushSubscription", "ContentEntity", "Suscripciones push notification (Web Push)"),
        ("Reseller", "ContentEntity", "Partners de reventa white-label"),
        ("ScheduledReport", "ContentEntity", "Reportes programados por tenant"),
        ("StylePreset", "ContentEntity", "Presets visuales por vertical"),
    ]
    add_styled_table(doc,
        ["Entidad", "Tipo", "Descripcion"],
        core_entities,
        [3.5, 4, 8.5]
    )

    doc.add_heading("9.3 Distribucion de Entidades por Modulo (Top 15)", level=2)
    entity_dist = [
        ("jaraba_agroconecta_core", "91", "Productos, lotes, trazabilidad, certificados, productores"),
        ("jaraba_comercio_conecta", "42", "Tiendas, productos, pedidos, ofertas, fidelizacion"),
        ("ecosistema_jaraba_core", "35", "Tenants, verticales, planes, compliance, badges"),
        ("jaraba_ai_agents", "20", "Agentes, herramientas, prompts, telemetria, benchmarks"),
        ("jaraba_candidate", "12", "Perfiles, CV, skills, experiencia, educacion"),
        ("jaraba_support", "12", "Tickets, SLA, articulos KB, respuestas, macros"),
        ("jaraba_copilot_v2", "10", "Interacciones, sesiones, feedback, contexto"),
        ("jaraba_analytics", "10", "Eventos, dashboards, reportes, segmentos"),
        ("jaraba_page_builder", "9", "Pages, templates, versiones, locks"),
        ("jaraba_site_builder", "9", "Sitios, paginas, navegacion, configuracion"),
        ("jaraba_mentoring", "9", "Mentores, mentees, sesiones, matching"),
        ("jaraba_governance", "9", "Politicas, clasificacion, retencion, lineage"),
        ("jaraba_customer_success", "9", "Health scores, playbooks, alertas, NPS"),
        ("jaraba_content_hub", "8", "Articulos, autores, categorias, comentarios"),
        ("jaraba_tenant_knowledge", "9", "KB articles, colecciones, embeddings"),
    ]
    add_styled_table(doc,
        ["Modulo", "Entidades", "Tipos principales"],
        entity_dist,
        [4.5, 1.5, 10]
    )

    doc.add_heading("9.4 Reglas de Entidades", level=2)
    entity_rules = [
        "ENTITY-FK-001: FKs intra-modulo = entity_reference. Cross-modulo = integer. tenant_id SIEMPRE entity_reference.",
        "AUDIT-CONS-001: TODA ContentEntity DEBE tener AccessControlHandler en anotacion.",
        "ENTITY-001: Toda entity con EntityOwnerTrait DEBE declarar EntityOwnerInterface + EntityChangedInterface.",
        "ENTITY-PREPROCESS-001: Toda entity con view mode DEBE tener template_preprocess_{type}() en .module.",
        "LABEL-NULLSAFE-001: 20+ entities no tienen 'label' en entity_keys. label() puede retornar NULL.",
        "FIELD-UI-SETTINGS-TAB-001: Toda entity con field_ui_base_route DEBE tener default local task tab.",
        "Views: Toda entity DEBE declarar views_data = Drupal\\views\\EntityViewsData en anotacion.",
    ]
    for rule in entity_rules:
        doc.add_paragraph(rule, style='List Bullet')

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════
    # 10. STACK DE IA
    # ═══════════════════════════════════════════════════════════════════
    doc.add_heading("10. Stack de Inteligencia Artificial", level=1)

    doc.add_paragraph(
        "Jaraba implementa un stack de IA \"Clase Mundial\" con 11 agentes Gen 2, "
        "streaming token-by-token, semantic cache, guardrails PII, MCP server, "
        "y tool use nativo. La IA no es un addon: esta integrada en el core del producto."
    )

    doc.add_heading("10.1 Arquitectura de Agentes Gen 2", level=2)
    doc.add_paragraph(
        "Todos los agentes extienden SmartBaseAgent con un constructor de 10 argumentos "
        "(6 core + 4 opcionales). Implementan doExecute() (no execute()). El metodo "
        "execute() maneja seleccion de experimentos A/B antes de delegar."
    )

    agents = [
        ("SmartMarketingAgent", "Marketing", "Optimizacion de campanas, copywriting, segmentacion"),
        ("StorytellingAgent", "Contenido", "Generacion de narrativas, storytelling de marca"),
        ("CustomerExperienceAgent", "CX", "Optimizacion de experiencia de cliente, journey mapping"),
        ("SupportAgent", "Soporte", "Clasificacion de tickets, drafts de respuesta, escalado"),
        ("ProducerCopilotAgent", "Productores", "Coaching de negocio para productores agroalimentarios"),
        ("SalesAgent", "Ventas", "Automatizacion de ventas, propuestas, follow-up"),
        ("MerchantCopilotAgent", "Comerciantes", "Guia para comerciantes locales, inventario, pricing"),
        ("SmartEmployabilityCopilot", "Empleo", "Coaching para buscadores de empleo, CV, entrevistas"),
        ("SmartLegalCopilot", "Legal", "Investigacion legal, asistencia en casos, plazos"),
        ("SmartContentWriter", "Escritura", "Creacion de contenido SEO, blog posts, copy"),
        ("LearningPathAgent", "Formacion", "Rutas de aprendizaje personalizadas por IA"),
    ]
    add_styled_table(doc,
        ["Agente", "Sector", "Funcion Principal"],
        agents,
        [4.5, 2.5, 9]
    )

    doc.add_heading("10.2 Model Routing (3 Tiers)", level=2)
    model_tiers = [
        ("Fast", "Haiku 4.5", "< 500ms", "Clasificacion, extraccion, filtros rapidos"),
        ("Balanced", "Sonnet 4.6", "1-3s", "Operaciones estandar, copilot chat, drafts"),
        ("Premium", "Opus 4.6", "3-8s", "Razonamiento complejo, analisis legal, estrategia"),
    ]
    add_styled_table(doc,
        ["Tier", "Modelo", "Latencia Tipica", "Casos de Uso"],
        model_tiers,
        [2, 3, 3, 8]
    )

    doc.add_heading("10.3 Capacidades Clave", level=2)

    ai_caps = [
        ("Streaming SSE", "Token-by-token via PHP Generator. Eventos: chunk, cached, done, error, thinking"),
        ("Tool Use Nativo", "callAiApiWithNativeTools() con ChatInput::setChatTools(). Max 5 iteraciones"),
        ("MCP Server", "POST /api/v1/mcp (JSON-RPC 2.0). Metodos: initialize, tools/list, tools/call, ping"),
        ("Semantic Cache", "Qdrant fuzzy matching antes de LLM. Reduce coste ~40% en queries similares"),
        ("PII Guardrails", "Deteccion bidireccional: DNI, NIE, IBAN ES, NIF/CIF, +34. 4 acciones: ALLOW/MODIFY/BLOCK/FLAG"),
        ("Long-Term Memory", "Per-agent en Qdrant: facts, preferences, interaction_summaries, corrections"),
        ("Multimodal", "Vision (GPT-4o), Audio (Whisper), Speech (TTS-1), Images (DALL-E 3)"),
        ("Distributed Tracing", "TraceContextService: trace_id + span_id + parent_span_id por request"),
        ("Self-Healing", "AutoDiagnosticService: tier override, provider rotation, throttle config via State API"),
        ("Brand Voice", "BrandVoiceProfile entity per-tenant. AIIdentityRule centralizado"),
        ("Observability", "AIObservabilityService: latency P95, error rates, cost estimation"),
        ("A/B Testing", "Prompt versioning + AgentBenchmarkService (LLM-as-Judge evaluation)"),
    ]
    add_styled_table(doc,
        ["Capacidad", "Detalle Tecnico"],
        ai_caps,
        [3, 13]
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════
    # 11. PAGE BUILDER
    # ═══════════════════════════════════════════════════════════════════
    doc.add_heading("11. Page Builder (GrapesJS)", level=1)

    doc.add_paragraph(
        "El page builder integra GrapesJS 5.7 directamente en el modulo jaraba_page_builder "
        "con 11 plugins custom, 202+ bloques en 24 categorias, y arquitectura dual "
        "(GrapesJS para edicion + Drupal behaviors para frontend publicado)."
    )

    doc.add_heading("11.1 Plugins Custom", level=2)
    plugins = [
        ("jaraba-canvas", "Core renderer del canvas con tema Jaraba"),
        ("jaraba-blocks", "Libreria de 202+ bloques organizados en 24 categorias"),
        ("jaraba-icons", "Insercion de SVG con 352+ iconos del sistema"),
        ("jaraba-ai", "Generacion de contenido por IA dentro del editor"),
        ("jaraba-seo", "Meta tags, Open Graph, Schema.org, auditor SEO"),
        ("jaraba-reviews", "Bloques de reviews y testimonios"),
        ("jaraba-forms", "Constructor de formularios integrado"),
        ("jaraba-table", "Tablas de datos responsive"),
        ("jaraba-media", "Libreria de imagenes y video del tenant"),
        ("jaraba-ecommerce", "Showcase de productos del marketplace"),
        ("jaraba-layout", "Sistema de grid responsive con 8 breakpoints"),
    ]
    add_styled_table(doc,
        ["Plugin", "Funcion"],
        plugins,
        [4, 12]
    )

    doc.add_heading("11.2 Categorias de Bloques", level=2)
    block_cats = [
        "Hero sections, Card grids, Call-to-action, Testimonios, FAQs (Schema.org), "
        "Features showcase, Pricing tables, Contact forms, Product showcase, Video embeds, "
        "Countdown timers, Social proof, Navigation menus, Footer layouts, Image galleries, "
        "Stats counters, Team members, Logos/Partners, Maps, Timeline, Comparison tables, "
        "Accordion/Tabs, Newsletter signup, Download sections"
    ]
    doc.add_paragraph(block_cats[0])

    doc.add_heading("11.3 Funcionalidades del Editor", level=2)
    editor_features = [
        "Editor full-viewport que bypasa el sistema de templates de pagina",
        "Preview responsive en 8 viewports (mobile, tablet, desktop, wide, ultra)",
        "55 templates verticalizados (templates x 5 verticales)",
        "Marketplace de templates compartidos entre tenants",
        "Editor multi-pagina con navegacion entre paginas del sitio",
        "Bloqueo de edicion concurrente (optimistic + pessimistic locking)",
        "Versionado con historial de cambios y rollback",
        "Validacion de accesibilidad WCAG integrada",
        "Asistente SEO con puntuacion y recomendaciones",
    ]
    for feat in editor_features:
        doc.add_paragraph(feat, style='List Bullet')

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════
    # 12. THEMING
    # ═══════════════════════════════════════════════════════════════════
    doc.add_heading("12. Sistema de Theming y Diseno", level=1)

    doc.add_heading("12.1 Arquitectura de Tema Unico", level=2)
    doc.add_paragraph(
        "ecosistema_jaraba_theme es el unico tema del ecosistema. No existe tema base "
        "ni temas hijos. Toda la personalizacion se logra via 5 capas de design tokens "
        "y CSS Custom Properties con prefijo --ej-*."
    )

    doc.add_heading("12.2 Sistema de 5 Capas de Tokens", level=2)
    token_layers = [
        ("1. SCSS Variables", "Fuente de verdad: _variables.scss", "Colores, tipografia, spacing, breakpoints"),
        ("2. CSS Custom Properties", "var(--ej-*, fallback)", "70+ variables configurables desde UI de tema"),
        ("3. Component Tokens", "Overrides por componente", "Botones, cards, headers, footers especificos"),
        ("4. Tenant Override", "TenantThemeConfig entity", "Cada tenant puede sobreescribir via UI o API"),
        ("5. Vertical Presets", "Combinaciones predefinidas", "Ej: agroconecta_nature_green, lex_corporate_blue"),
    ]
    add_styled_table(doc,
        ["Capa", "Mecanismo", "Alcance"],
        token_layers,
        [3.5, 4, 8.5]
    )

    doc.add_heading("12.3 Colores de Marca", level=2)
    colors = [
        ("--ej-corporativo-primario", "#233D63", "Azul corporativo. Header, links, CTAs primarios"),
        ("--ej-impulso", "#FF8C42", "Naranja impulso. Badges, highlights, CTAs secundarios"),
        ("--ej-innovacion", "#00A9A5", "Verde/teal innovacion. Success, features, iconos"),
    ]
    add_styled_table(doc,
        ["Variable CSS", "Valor", "Uso"],
        colors,
        [4.5, 2.5, 9]
    )

    doc.add_heading("12.4 Arquitectura SCSS", level=2)
    scss_stats = [
        ("Archivos SCSS totales", "102"),
        ("Partials (_*.scss)", "19"),
        ("Bundles (bundles/*.scss)", "8"),
        ("Components (components/_*.scss)", "62"),
        ("Features (features/_*.scss)", "3 (back-to-top, dark-mode, promo-banner)"),
        ("Routes (routes/*.scss)", "13+"),
        ("Compilador", "Dart Sass (moderno, @use, NO @import)"),
        ("Build", "npm run build desde ecosistema_jaraba_theme/"),
    ]
    add_styled_table(doc,
        ["Metrica", "Valor"],
        scss_stats,
        [5.5, 10.5]
    )

    doc.add_heading("12.5 Sistema de Iconos", level=2)
    doc.add_paragraph(
        "352 iconos SVG organizados en 6 categorias primarias (actions, fiscal, media, micro, ui, users) "
        "y 7 categorias bridge (symlinks). Funcion Twig: jaraba_icon('category', 'name', {variant, color, size}). "
        "Variantes: outline (default), outline-bold, filled, duotone. "
        "Colores exclusivos de paleta Jaraba: azul-corporativo, naranja-impulso, verde-innovacion, white, neutral."
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════
    # 13. SEO Y GEO
    # ═══════════════════════════════════════════════════════════════════
    doc.add_heading("13. SEO, GEO y Descubribilidad por IA", level=1)

    doc.add_heading("13.1 SEO Tradicional", level=2)
    seo_features = [
        ("Meta Tags", "Metatag module + SeoService custom: title, description, canonical, robots"),
        ("Open Graph", "og:title, og:description, og:image, og:type para cada pagina"),
        ("Twitter Cards", "twitter:card, twitter:title, twitter:description, twitter:image"),
        ("Schema.org JSON-LD", "SoftwareApplication, FAQPage, BreadcrumbList, Offer por tier"),
        ("Sitemap XML", "simple_sitemap con prioridades por tipo de contenido y vertical"),
        ("URLs Semanticas", "Pathauto + redirect module. Slugs limpios por entidad"),
        ("Breadcrumbs", "BreadcrumbList Schema.org + navegacion visual jerarquica"),
        ("hreflang", "Multi-idioma (ES + EN + PT-BR) con alternate links"),
        ("Core Web Vitals", "LCP: fetchpriority=high. CLS: dimensiones fijas. INP: handlers debounced"),
    ]
    add_styled_table(doc,
        ["Componente", "Implementacion"],
        seo_features,
        [3.5, 12.5]
    )

    doc.add_heading("13.2 GEO: Generative Engine Optimization", level=2)
    doc.add_paragraph(
        "GEO es la optimizacion para que el contenido sea descubierto y recomendado "
        "por motores de IA generativa (ChatGPT, Gemini, Claude, Perplexity). "
        "Jaraba implementa las siguientes estrategias:"
    )
    geo_features = [
        "llms.txt: Endpoint /llms.txt que expone estructura del sitio optimizada para crawlers IA",
        "Schema.org enriquecido: JSON-LD con datos estructurados que los LLMs consumen directamente",
        "Contenido factual y citeable: Estructura que facilita la extraccion de facts por IA",
        "FAQ semantica: Preguntas frecuentes con Schema.org FAQPage para featured snippets IA",
        "Reviews y calificaciones: AggregateRating Schema.org para social proof en respuestas IA",
        "Datos estructurados de productos: Pricing, availability, offers para comparaciones IA",
    ]
    for feat in geo_features:
        doc.add_paragraph(feat, style='List Bullet')

    doc.add_heading("13.3 Estrategia de Contenido para Descubribilidad", level=2)
    doc.add_paragraph(
        "El Content Hub genera articulos optimizados para SEO y GEO simultaneamente. "
        "SmartContentWriterAgent produce contenido que combina keywords tradicionales "
        "con estructura factual que los LLMs prefieren citar. Cada articulo incluye "
        "meta tags, Schema.org, Open Graph y estructura de headings optimizada."
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════
    # 14. UX
    # ═══════════════════════════════════════════════════════════════════
    doc.add_heading("14. Experiencia de Usuario (UX)", level=1)

    doc.add_heading("14.1 Zero Region Pattern", level=2)
    doc.add_paragraph(
        "Las paginas frontend usan el Zero Region Pattern: cada ruta tiene un template "
        "page--{ruta}.html.twig con layout limpio. Se usa {{ clean_content }} en vez de "
        "{{ page.content }} para extraer solo el bloque principal. Los datos se inyectan "
        "via drupalSettings en hook_preprocess_page(), nunca desde el controller."
    )

    doc.add_heading("14.2 Slide-Panel Forms", level=2)
    doc.add_paragraph(
        "TODA accion crear/editar/ver en frontend se abre en slide-panel lateral (no navega fuera). "
        "Usa renderPlain() (no render()) para evitar BigPipe placeholders. "
        "Deteccion: isXmlHttpRequest() && !_wrapper_format."
    )

    doc.add_heading("14.3 Parciales Reutilizables (65+)", level=2)
    partials = [
        ("_header.html.twig", "Header con 3 variantes: classic, minimal, transparent"),
        ("_footer.html.twig", "Footer con 3 layouts: mega, standard, split"),
        ("_avatar-nav.html.twig", "Navegacion contextual para 10 tipos de avatar/rol"),
        ("_copilot-fab.html.twig", "Boton flotante del copilot IA"),
        ("_command-bar.html.twig", "Paleta de comandos tipo CMD+K"),
        ("_bottom-nav.html.twig", "Navegacion inferior mobile"),
        ("_skeleton.html.twig", "Loading states con skeleton screens"),
        ("_empty-state.html.twig", "Estados vacios con ilustracion y CTA"),
    ]
    add_styled_table(doc,
        ["Parcial", "Funcion"],
        partials,
        [4.5, 11.5]
    )

    doc.add_heading("14.4 Responsive Design", level=2)
    doc.add_paragraph(
        "Mobile-first con SCSS media queries ascendentes. 8 breakpoints definidos: "
        "mobile (320px), mobile-lg (480px), tablet (768px), tablet-lg (1024px), "
        "desktop (1280px), desktop-lg (1440px), wide (1920px), ultra (2560px). "
        "Flexbox + CSS Grid. Lenis smooth scroll. Core Web Vitals optimizados."
    )

    doc.add_heading("14.5 Accesibilidad", level=2)
    doc.add_paragraph(
        "WCAG 2.1 AA como objetivo: aria-labels en todos los interactivos, "
        "headings jerarquicos, focus visible, contraste verificado, "
        "soporte de screen readers, responsive motion, touch targets minimos."
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════
    # 15. SEGURIDAD
    # ═══════════════════════════════════════════════════════════════════
    doc.add_heading("15. Seguridad, Compliance y Gobernanza", level=1)

    doc.add_heading("15.1 Gestion de Secrets", level=2)
    doc.add_paragraph(
        "SECRET-MGMT-001: 3 capas de gestion de secretos. "
        "Capa 1: config/sync/ en Git (valores vacios). "
        "Capa 2: settings.secrets.php con getenv(). "
        "Capa 3: Variables de entorno (local .env, produccion IONOS panel). "
        "14 configuraciones protegidas: OAuth (3 proveedores), SMTP, reCAPTCHA, Stripe."
    )

    doc.add_heading("15.2 Control de Acceso", level=2)
    access_features = [
        "39+ AccessControlHandlers distribuidos en 21 modulos",
        "TENANT-ISOLATION-ACCESS-001: Verificacion de tenant match para update/delete",
        "ACCESS-STRICT-001: Comparaciones (int) === (int) para ownership (previene type juggling)",
        "CSRF-API-001: API routes via fetch() usan X-CSRF-Request-Header-Token",
        "Webhooks con HMAC + hash_equals() (nunca query string tokens)",
        "reCAPTCHA v3 en formularios publicos",
    ]
    for feat in access_features:
        doc.add_paragraph(feat, style='List Bullet')

    doc.add_heading("15.3 Compliance Stack", level=2)
    compliance = [
        ("GDPR / LOPD-GDD", "jaraba_privacy", "DPA, consentimiento, derecho al olvido, exportacion datos"),
        ("SOC 2 / ISO 27001", "jaraba_security_compliance", "25+ controles, audit trail, scoring A-F"),
        ("ENS (Esquema Nacional)", "jaraba_security_compliance", "Framework espanol de seguridad"),
        ("EU AI Act", "jaraba_ai_agents", "Audit trail de decisiones IA, transparencia"),
        ("Facturacion Electronica", "jaraba_facturae + jaraba_verifactu", "Facturae 3.2.2, VeriFactu RD 1007/2023"),
        ("CAN-SPAM", "jaraba_email", "Compliance en 28 templates MJML"),
    ]
    add_styled_table(doc,
        ["Marco", "Modulo", "Implementacion"],
        compliance,
        [3, 4.5, 8.5]
    )

    doc.add_heading("15.4 Auditoria y Monitoring", level=2)
    doc.add_paragraph(
        "AuditLog entity (append-only, inmutable) con 25+ controles tracked. "
        "ComplianceAggregatorService calcula 9 KPIs con score 0-100 y grados A-F. "
        "CompliancePanelController en /admin/jaraba/compliance con auto-refresh 60s. "
        "Security scan diario via GitHub Actions (Trivy + OWASP ZAP + audits)."
    )

    doc.add_heading("15.5 Guardrails de IA", level=2)
    doc.add_paragraph(
        "AIGuardrailsService implementa proteccion bidireccional (input + output): "
        "deteccion de PII (DNI, NIE, IBAN ES, NIF/CIF, +34), prevencion de jailbreak, "
        "mascarado de output, con 4 niveles de accion (ALLOW, MODIFY, BLOCK, FLAG). "
        "AIIdentityRule centraliza la identidad de marca en todos los agentes."
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════
    # 16. INFRAESTRUCTURA Y DEVOPS
    # ═══════════════════════════════════════════════════════════════════
    doc.add_heading("16. Infraestructura y DevOps", level=1)

    doc.add_heading("16.1 Entornos", level=2)
    envs = [
        ("Desarrollo Local", "Lando (Docker)", "MariaDB + Redis + Qdrant + Tika + Mailhog",
         "https://jaraba-saas.lndo.site/"),
        ("Produccion", "IONOS Dedicated L-16", "128GB RAM, AMD EPYC, NVMe SSD",
         "Deploy automatico via GitHub Actions"),
    ]
    add_styled_table(doc,
        ["Entorno", "Plataforma", "Recursos", "Acceso"],
        envs,
        [3, 3, 5, 5]
    )

    doc.add_heading("16.2 Pipeline CI/CD (8 Workflows)", level=2)
    workflows = [
        ("ci.yml", "Push/PR", "Lint (PHPCS), PHPStan L6, Unit tests, Kernel tests (MariaDB), SCSS build"),
        ("deploy.yml", "Push main", "Backup DB -> SSH deploy -> drush updatedb -> config:import -> cache:rebuild -> smoke test"),
        ("security-scan.yml", "Diario", "Trivy filesystem, Composer audit, npm audit, OWASP ZAP"),
        ("deploy-production.yml", "Tag release", "Blue-Green container deploy con rollback"),
        ("deploy-staging.yml", "Branch staging", "Deploy a entorno de pre-produccion"),
        ("daily-backup.yml", "Cron 03:00 UTC", "Backup automatico de base de datos"),
        ("fitness-functions.yml", "Push", "Validacion arquitectural (dependencias, metricas)"),
        ("verify-backups.yml", "Semanal", "Verificacion de integridad de backups"),
    ]
    add_styled_table(doc,
        ["Workflow", "Trigger", "Acciones"],
        workflows,
        [3.5, 2.5, 10]
    )

    doc.add_heading("16.3 Monitoring Stack", level=2)
    monitoring = [
        ("Prometheus", "9090", "Scraping de metricas del servidor y aplicacion"),
        ("Grafana", "3001", "Dashboards visuales con alertas"),
        ("Loki", "3100", "Agregacion de logs centralizada"),
        ("AlertManager", "9093", "14 reglas de alerta -> Slack/email"),
    ]
    add_styled_table(doc,
        ["Herramienta", "Puerto", "Funcion"],
        monitoring,
        [3, 2, 11]
    )

    doc.add_heading("16.4 Lando Tooling (Desarrollo)", level=2)
    tooling = [
        "lando drush — CLI de Drupal (cache, config, migraciones)",
        "lando composer — Gestion de dependencias PHP",
        "lando redis-cli / redis-status — Administracion de cache Redis",
        "lando qdrant-status / qdrant-health — Estado del vector DB",
        "lando tika-test — Verificacion de Apache Tika",
        "lando ai-health — Health check completo del stack de IA",
        "lando validate — Validacion arquitectural completa",
        "lando validate-fast — Validacion rapida pre-commit",
        "lando self-heal — Herramientas de auto-recovery",
    ]
    for tool in tooling:
        doc.add_paragraph(tool, style='List Bullet')

    doc.add_heading("16.5 Testing", level=2)
    doc.add_paragraph(
        "453 tests automatizados (Unit + Kernel) con 80% coverage minimo en modulos jaraba_*. "
        "phpunit.xml en raiz. CI ejecuta ambas suites con MariaDB 10.11 service container. "
        "Suites: Unit, Kernel, Functional, PromptRegression."
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════
    # 17. PLANES SAAS
    # ═══════════════════════════════════════════════════════════════════
    doc.add_heading("17. Planes SaaS y Estructura de Precios", level=1)

    doc.add_heading("17.1 Arquitectura de Pricing", level=2)
    doc.add_paragraph(
        "El sistema de precios utiliza una cascada de resolucion de 3 niveles gestionada "
        "por MetaSitePricingService -> PlanResolverService:"
    )
    cascade = [
        ("1. Especifico", "{vertical}_{tier}", "Ej: agroconecta_professional — precios de AgroConecta Pro"),
        ("2. Default", "_default_{tier}", "Ej: _default_professional — precios genericos Pro"),
        ("3. Fallback", "Hardcoded en servicio", "Precios de emergencia si no hay ConfigEntity"),
    ]
    add_styled_table(doc,
        ["Nivel", "Patron", "Ejemplo"],
        cascade,
        [2.5, 3.5, 10]
    )

    doc.add_heading("17.2 Planes por Tier", level=2)
    tier_details = [
        ("Starter", "0 EUR/mes", "0 EUR/ano",
         "Plan gratuito para siempre. 1 usuario, funcionalidades basicas, sin IA, ideal para evaluacion."),
        ("Professional", "29 EUR/mes", "290 EUR/ano (17% dto)",
         "IA incluida, hasta 5 usuarios, firma digital FNMT, analytics avanzado, soporte prioritario."),
        ("Enterprise", "99 EUR/mes", "990 EUR/ano (17% dto)",
         "Usuarios ilimitados, white-label, blockchain, API completa, soporte premium dedicado."),
    ]
    add_styled_table(doc,
        ["Plan", "Mensual", "Anual", "Incluye"],
        tier_details,
        [2.5, 2.5, 3, 8]
    )

    doc.add_heading("17.3 Verticalizacion de Precios", level=2)
    doc.add_paragraph(
        "Cada vertical puede definir precios, limites y features diferentes mediante "
        "SaasPlan (ContentEntity) asociado a un Vertical + SaasPlanTier. "
        "La pagina /planes muestra precios genericos. "
        "La pagina /planes/{vertical} muestra precios especificos del vertical. "
        "El PricingController resuelve los datos via MetaSitePricingService."
    )

    doc.add_heading("17.4 Entidades del Sistema de Pricing", level=2)
    pricing_entities = [
        ("SaasPlanTier", "ConfigEntity", "plan_tier", "starter, professional, enterprise — aliases, Stripe IDs"),
        ("SaasPlan", "ContentEntity", "saas_plan", "Precios EUR por vertical+tier, peso, limites"),
        ("Vertical", "ContentEntity", "vertical", "machine_name, status, descripcion del sector"),
        ("Tenant", "ContentEntity", "tenant", "plan_id FK a SaasPlan, Stripe customer, dominio"),
    ]
    add_styled_table(doc,
        ["Entidad", "Tipo", "Machine Name", "Funcion"],
        pricing_entities,
        [3, 2.5, 2.5, 8]
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════
    # 18. ANALISIS FINANCIERO
    # ═══════════════════════════════════════════════════════════════════
    doc.add_heading("18. Analisis Financiero y Proyecciones", level=1)

    doc.add_heading("18.1 Contexto de Mercado", level=2)
    doc.add_paragraph(
        "El mercado SaaS en Espana alcanza ~8.160M EUR con CAGR del 19% (2024-2029). "
        "El SaaS vertical crece 2x mas rapido que el horizontal. "
        "Kit Digital ha emitido 880.000+ bonos (2022-2025). "
        "El 26% de plataformas SaaS ya tienen IA embebida (2025)."
    )

    doc.add_heading("18.2 Proyecciones Financieras (36 meses, escenario conservador)", level=2)
    financials = [
        ("Tenants Activos", "80", "200", "400"),
        ("ARPU/Mes", "65 EUR", "90 EUR", "120 EUR"),
        ("MRR", "5.200 EUR", "18.000 EUR", "48.000 EUR"),
        ("ARR", "62.400 EUR", "216.000 EUR", "576.000 EUR"),
        ("Churn Mensual", "6%", "4%", "3%"),
        ("NRR", "95%", "105%", "112%"),
        ("CAC", "250 EUR", "400 EUR", "500 EUR"),
        ("LTV (36m)", "975 EUR", "2.070 EUR", "3.840 EUR"),
        ("LTV:CAC", "3,9:1", "5,2:1", "7,7:1"),
        ("Margen Bruto", "72%", "78%", "82%"),
        ("Resultado Neto", "-132.000 EUR", "-63.000 EUR", "+120.000 EUR"),
    ]
    add_styled_table(doc,
        ["Metrica", "Ano 1 (M12)", "Ano 2 (M24)", "Ano 3 (M36)"],
        financials,
        [3.5, 4, 4, 4]
    )

    doc.add_heading("18.3 Estructura de Costes", level=2)
    costs = [
        ("Desarrollo", "108.000 EUR", "72.000 EUR", "54.000 EUR"),
        ("Infraestructura", "12.000 EUR", "18.000 EUR", "24.000 EUR"),
        ("Salarios", "36.000 EUR", "96.000 EUR", "168.000 EUR"),
        ("Marketing", "6.000 EUR", "24.000 EUR", "48.000 EUR"),
        ("Operaciones", "12.000 EUR", "15.000 EUR", "18.000 EUR"),
        ("TOTAL", "174.000 EUR", "225.000 EUR", "312.000 EUR"),
    ]
    add_styled_table(doc,
        ["Concepto", "Ano 1", "Ano 2", "Ano 3"],
        costs,
        [3.5, 4, 4, 4]
    )

    p = doc.add_paragraph()
    run = p.add_run("Break-even estimado: ")
    run.bold = True
    run.font.color.rgb = BLUE
    p.add_run("Mes 20-22 (Q4 2027)")

    doc.add_heading("18.4 KPIs de Decision", level=2)
    kpis = [
        ("Mes 3", "3 pilotos activos", "Go/No-Go Fase 1"),
        ("Mes 6", "20 tenants, NPS >30, activation >40%", "Go/No-Go Fase 2 (comercial)"),
        ("Mes 12", "50 tenants, churn <5%, MRR >4K EUR", "Go/No-Go expansion nacional"),
        ("Mes 18", "NRR >100%, LTV:CAC >3:1", "Evaluar fundraising"),
    ]
    add_styled_table(doc,
        ["Hito", "Criterio", "Decision"],
        kpis,
        [2.5, 7, 6.5]
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════
    # 19. GTM
    # ═══════════════════════════════════════════════════════════════════
    doc.add_heading("19. Estrategia Go-To-Market", level=1)

    doc.add_heading("19.1 Estrategia \"Submarinos con Periscopio\"", level=2)
    doc.add_paragraph(
        "Cada vertical se lanza como producto independiente con sub-marca propia. "
        "La plataforma integrada solo se revela cuando anade valor al cliente. "
        "Esto evita la paralisis por analisis del comprador ante 10 verticales simultaneos "
        "y permite posicionamiento de nicho en cada sector."
    )

    doc.add_heading("19.2 Fases de Lanzamiento", level=2)
    phases = [
        ("Fase 0: Piloto", "Meses 1-3", "3-5 entidades PIIL",
         "Empleabilidad", "Activation >40%, D30 >25%, NPS >30"),
        ("Fase 1: Institucional", "Meses 3-6", "20 tenants, 8 provincias",
         "Empleabilidad + Emprendimiento", "NPS >40, Churn <5%"),
        ("Fase 2: Comercial", "Meses 6-12", "50+ tenants, 4 verticales",
         "+ ComercioConecta + AgroConecta", "NRR >100%, MRR 4K+ EUR"),
        ("Fase 3-4: Expansion", "Mes 12+", "150+ tenants",
         "Nacional + Internacional", "LTV:CAC >3:1"),
    ]
    add_styled_table(doc,
        ["Fase", "Timeline", "Objetivo", "Verticales", "KPIs"],
        phases,
        [2.5, 2.5, 3, 3.5, 4.5]
    )

    doc.add_heading("19.3 Canal Institucional (Ventaja Competitiva)", level=2)
    doc.add_paragraph(
        "Plataforma de Ecosistemas Digitales S.L. cuenta con 30+ anos de capital relacional "
        "con entidades institucionales (PIIL, SAE, GDR, ayuntamientos, diputaciones). "
        "Este canal B2G es el principal moat competitivo: las instituciones validan credibilidad, "
        "generan case studies, y actuan como multiplicador de adopcion (cada PIIL tiene 100-500 "
        "participantes potenciales que necesitan la plataforma)."
    )

    doc.add_heading("19.4 Argumentos de Venta por Vertical", level=2)
    sales_args = [
        ("Empleabilidad (PIIL)", "Automatiza reporting SEPE, LMS integrado, 79 EUR/mes vs 3K+ EUR enterprise"),
        ("AgroConecta", "Venta directa (8% comision vs 25% marketplaces), trazabilidad QR, Kit Digital eligible"),
        ("ComercioConecta", "Tienda online en 24h, 29 EUR/mes (< Shopify 36 EUR), QR dinamicos, fidelizacion"),
        ("ServiciosConecta", "Todo-en-uno (agenda + firma + presupuestos), 79 EUR/mes, no 4 SaaS separados"),
        ("Instituciones", "White-label, multi-programa, medicion de impacto, GDPR + accesibilidad"),
    ]
    add_styled_table(doc,
        ["Segmento", "Argumento Principal"],
        sales_args,
        [3.5, 12.5]
    )

    doc.add_heading("19.5 Principios de Ejecucion", level=2)
    exec_principles = [
        "FOCO ANTES QUE AMPLITUD: Lanzar solo Empleabilidad; los otros 9 esperan PMF",
        "INSTITUCIONAL PRIMERO: Aprovechar 30+ anos de capital relacional como moat",
        "SUBMARINOS CON PERISCOPIO: Cada vertical es sub-marca independiente",
        "METRICAS, NO OPINIONES: Activation >40%, D30 >25%, NPS >40, churn <5%",
        "VELOCIDAD DE ITERACION: Ciclos semanales feedback-implementar-medir",
        "HONESTIDAD COMERCIAL: Sin humo; verificar claims; credibilidad institucional",
        "BOOTSTRAP HASTA PMF: Sin funding externo hasta 50 tenants, NRR >100%, churn <5%",
    ]
    for principle in exec_principles:
        doc.add_paragraph(principle, style='List Bullet')

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════
    # 20. ROADMAP
    # ═══════════════════════════════════════════════════════════════════
    doc.add_heading("20. Roadmap y Madurez del Producto", level=1)

    doc.add_heading("20.1 Nivel de Madurez Actual", level=2)
    add_kpi_box(doc, "Madurez", "5.0 / 5.0", "Production-Ready, Resilience & Compliance Certified")

    doc.add_heading("20.2 Completado Recientemente (Q1 2026)", level=2)
    completed = [
        "Elevacion de landing pages (patron 3 niveles para Empleabilidad)",
        "Pipeline Claude Code DX (14 ficheros, 30 tests funcionales)",
        "Meta-sites 3 idiomas (ES + EN + PT-BR)",
        "Remediacion de secrets (SECRET-MGMT-001 completo)",
        "Stack de analytics completo (GTM + A/B + Heatmap + Tracking)",
        "Auditoria IA completa (30/30 items certificados)",
        "Stack IA clase mundial (33 items: 23 FIX + 10 GAP resueltos)",
        "Streaming real token-by-token (SSE)",
        "MCP server + native function calling",
        "Verticalizacion completa de precios (7 verticales con SaasPlan entities)",
        "11 agentes Gen 2 operativos",
        "Tenant remediation completada",
        "Core Web Vitals optimizados",
    ]
    for item in completed:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading("20.3 Prioridades Inmediatas (Q2 2026)", level=2)
    priorities = [
        "Go-Live Fase 0 (Piloto Institucional con 3-5 PIIL)",
        "Activacion de 38 modulos inactivos restantes",
        "Configuracion Stripe Connect produccion (plan tiers + webhooks)",
        "Onboarding flow completo (registro -> seleccion plan -> setup -> welcome)",
        "Performance optimization (p95 latencia < 500ms)",
        "Documentacion de API para integradores",
        "Kit Digital: homologacion como solucion elegible",
    ]
    for item in priorities:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading("20.4 Horizonte Medio (Q3-Q4 2026)", level=2)
    medium_term = [
        "Expansion a 4 verticales operativos",
        "50 tenants activos target",
        "Internacionalizacion (Portugal, LATAM piloto)",
        "Mobile app (backend via jaraba_mobile)",
        "Marketplace de templates y plugins",
        "Advanced analytics con ML predictivo",
    ]
    for item in medium_term:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════
    # ANEXO A: LISTADO COMPLETO DE MODULOS
    # ═══════════════════════════════════════════════════════════════════
    doc.add_heading("Anexo A: Listado Completo de 95 Modulos Custom", level=1)

    all_modules = [
        "ai_provider_google_gemini", "ecosistema_jaraba_core",
        "jaraba_ab_testing", "jaraba_addons", "jaraba_ads",
        "jaraba_agent_flows", "jaraba_agent_market", "jaraba_agents",
        "jaraba_agroconecta_core", "jaraba_ai_agents", "jaraba_ambient_ux",
        "jaraba_analytics", "jaraba_andalucia_ei", "jaraba_billing",
        "jaraba_blog", "jaraba_business_tools", "jaraba_candidate",
        "jaraba_comercio_conecta", "jaraba_commerce", "jaraba_connector_sdk",
        "jaraba_content_hub", "jaraba_copilot_v2", "jaraba_credentials",
        "jaraba_crm", "jaraba_customer_success", "jaraba_diagnostic",
        "jaraba_dr", "jaraba_einvoice_b2b", "jaraba_email",
        "jaraba_events", "jaraba_facturae", "jaraba_foc",
        "jaraba_funding", "jaraba_geo", "jaraba_governance",
        "jaraba_groups", "jaraba_heatmap", "jaraba_i18n",
        "jaraba_identity", "jaraba_insights_hub", "jaraba_institutional",
        "jaraba_integrations", "jaraba_interactive", "jaraba_job_board",
        "jaraba_journey", "jaraba_legal", "jaraba_legal_billing",
        "jaraba_legal_calendar", "jaraba_legal_cases", "jaraba_legal_intelligence",
        "jaraba_legal_knowledge", "jaraba_legal_lexnet", "jaraba_legal_templates",
        "jaraba_legal_vault", "jaraba_lms", "jaraba_matching",
        "jaraba_mentoring", "jaraba_messaging", "jaraba_mobile",
        "jaraba_multiregion", "jaraba_notifications", "jaraba_onboarding",
        "jaraba_page_builder", "jaraba_paths", "jaraba_performance",
        "jaraba_pixels", "jaraba_predictive", "jaraba_privacy",
        "jaraba_pwa", "jaraba_rag", "jaraba_referral",
        "jaraba_resources", "jaraba_security_compliance", "jaraba_self_discovery",
        "jaraba_sepe_teleformacion", "jaraba_servicios_conecta", "jaraba_site_builder",
        "jaraba_skills", "jaraba_sla", "jaraba_social",
        "jaraba_social_commerce", "jaraba_sso", "jaraba_success_cases",
        "jaraba_support", "jaraba_tenant_export", "jaraba_tenant_knowledge",
        "jaraba_theming", "jaraba_training", "jaraba_usage_billing",
        "jaraba_verifactu", "jaraba_whitelabel", "jaraba_workflows",
        "jaraba_zkp",
    ]

    # Display as numbered list in 2 columns via table
    half = (len(all_modules) + 1) // 2
    mod_rows = []
    for i in range(half):
        left = f"{i+1}. {all_modules[i]}"
        right = f"{i+half+1}. {all_modules[i+half]}" if i + half < len(all_modules) else ""
        mod_rows.append((left, right))

    add_styled_table(doc, ["#  Modulo", "#  Modulo"], mod_rows, [8, 8])

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════
    # ANEXO B: MATRIZ DE ENTIDADES
    # ═══════════════════════════════════════════════════════════════════
    doc.add_heading("Anexo B: Matriz de Entidades por Modulo (Top 20)", level=1)

    entity_matrix = [
        ("jaraba_agroconecta_core", "91", "Productor, Producto, Lote, Certificado, Pedido, Envio, Trazabilidad..."),
        ("jaraba_comercio_conecta", "42", "Tienda, ProductoLocal, Oferta, QRDinamico, ProgramaFidelidad..."),
        ("ecosistema_jaraba_core", "35", "Tenant, Vertical, SaasPlan, Badge, Feature, AuditLog, Reseller..."),
        ("jaraba_ai_agents", "20", "AIAgent, ToolDefinition, PromptTemplate, BrandVoiceProfile, AIUsageLog..."),
        ("jaraba_candidate", "12", "CandidateProfile, Skill, Experience, Education, Application..."),
        ("jaraba_support", "12", "SupportTicket, SupportArticle, SLAPolicy, TicketResponse, Macro..."),
        ("jaraba_copilot_v2", "10", "CopilotInteraction, CopilotSession, CopilotFeedback, OsterwalderExperiment..."),
        ("jaraba_analytics", "10", "AnalyticsEvent, Dashboard, Report, Segment, Funnel..."),
        ("jaraba_page_builder", "9", "PageContent, PageTemplate, PageVersion, EditLock, QuotaUsage..."),
        ("jaraba_site_builder", "9", "Site, SitePage, SiteNavigation, SiteConfig..."),
        ("jaraba_mentoring", "9", "Mentor, Mentee, MentoringSession, MentoringMatch..."),
        ("jaraba_governance", "9", "DataPolicy, DataClassification, RetentionRule, DataLineage..."),
        ("jaraba_customer_success", "9", "HealthScore, RetentionPlaybook, ChurnAlert, NPSSurvey..."),
        ("jaraba_tenant_knowledge", "9", "KBArticle, KBCollection, KBEmbedding, KBFeedback..."),
        ("jaraba_content_hub", "8", "ContentArticle, ContentAuthor, ContentCategory, ContentComment..."),
        ("jaraba_billing", "5", "Subscription, Invoice, UsageRecord, ImpactCredit, Trial..."),
        ("jaraba_legal_cases", "5", "LegalCase, CaseParty, CaseDocument, CaseEvent, CaseDeadline"),
        ("jaraba_crm", "5", "Contact, Company, Deal, Activity, Pipeline..."),
        ("jaraba_credentials", "4", "Credential, CredentialStack, CredentialTemplate, BadgeClass"),
        ("jaraba_legal_vault", "4", "VaultDocument, VaultFolder, VaultShare, VaultAudit"),
    ]
    add_styled_table(doc,
        ["Modulo", "Total", "Entidades Principales"],
        entity_matrix,
        [4, 1.5, 10.5]
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════
    # ANEXO C: REGLAS Y DIRECTRICES
    # ═══════════════════════════════════════════════════════════════════
    doc.add_heading("Anexo C: Reglas y Directrices del Proyecto (Seleccion)", level=1)

    doc.add_paragraph(
        "El proyecto cuenta con 178+ especificaciones tecnicas documentadas. "
        "A continuacion se listan las mas relevantes para directivos y tecnicos."
    )

    rules_table = [
        ("TENANT-001", "Toda query filtra por tenant_id", "Critica"),
        ("TENANT-BRIDGE-001", "Usar TenantBridgeService para Tenant<->Group", "Critica"),
        ("SECRET-MGMT-001", "3 capas de secrets, nunca en config/sync", "Critica"),
        ("PREMIUM-FORMS-PATTERN-001", "Toda entity form extiende PremiumEntityFormBase", "Alta"),
        ("AUDIT-CONS-001", "Toda entity tiene AccessControlHandler", "Critica"),
        ("AGENT-GEN2-PATTERN-001", "Agentes extienden SmartBaseAgent, override doExecute()", "Alta"),
        ("MODEL-ROUTING-CONFIG-001", "3 tiers IA en YAML (fast/balanced/premium)", "Alta"),
        ("AI-GUARDRAILS-PII-001", "Deteccion PII bidireccional", "Critica"),
        ("CSS-VAR-ALL-COLORS-001", "Todo color via var(--ej-*, fallback)", "Alta"),
        ("ICON-CONVENTION-001", "jaraba_icon() con paleta de colores restringida", "Media"),
        ("ROUTE-LANGPREFIX-001", "URLs via Url::fromRoute(), nunca hardcoded", "Alta"),
        ("UPDATE-HOOK-CATCH-001", "catch(\\Throwable) en hooks, no \\Exception", "Critica"),
        ("SCSS-COMPILE-VERIFY-001", "Verificar timestamp CSS > SCSS tras edicion", "Alta"),
        ("DOC-GUARD-001", "Never overwrite master docs, solo Edit incremental", "Alta"),
        ("ZERO-REGION-001", "drupalSettings via preprocess, no controller", "Alta"),
        ("SLIDE-PANEL-RENDER-001", "renderPlain() para slide-panel forms", "Alta"),
        ("FORM-CACHE-001", "Nunca setCached(TRUE) incondicional", "Alta"),
        ("ACCESS-STRICT-001", "Comparaciones (int)===(int) para ownership", "Critica"),
        ("CSRF-API-001", "X-CSRF-Request-Header-Token para API fetch", "Critica"),
        ("STREAMING-GAP01-001", "SSE via PHP Generator, eventos tipados", "Alta"),
    ]
    add_styled_table(doc,
        ["Regla", "Descripcion", "Prioridad"],
        rules_table,
        [4.5, 9, 2.5]
    )

    # ═══════════════════════════════════════════════════════════════════
    # FINAL PAGE
    # ═══════════════════════════════════════════════════════════════════
    doc.add_page_break()
    for _ in range(8):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Jaraba Impact Platform")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = BLUE

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Plataforma de Ecosistemas Digitales S.L.")
    run.font.size = Pt(12)
    run.font.color.rgb = GRAY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Documento Comprensivo v1.0.0 — Marzo 2026")
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CONFIDENCIAL — Uso Interno")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = ORANGE

    # ─── Save ───
    output_path = os.path.join(OUTPUT_DIR, FILENAME)
    doc.save(output_path)
    print(f"Documento generado: {output_path}")
    print(f"Tamano: {os.path.getsize(output_path) / 1024:.1f} KB")
    return output_path


if __name__ == "__main__":
    build_document()
